
import json
import datetime
from django.http import JsonResponse, HttpResponse
from django.shortcuts import render, get_object_or_404
from django.views.decorators.csrf import csrf_exempt
from .registry import run_action
from scanner.utils.normalize_url import normalize_url
from scanner.models import ScanReport
from scanner.utils.pdf_generator import generate_security_report_pdf
from django.conf import settings
from scanner.utils.docx_generator import generate_security_report_docx
from scanner.utils.extract_page_title import get_page_title, get_company_name_from_title
import logging
logger = logging.getLogger(__name__)

from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def scan_page(request):
    return render(request, "scanner/scan.html")

from asgiref.sync import sync_to_async

async def scan_api(request):
    if request.method != "POST":
        return JsonResponse({"status": "error", "error": "POST only"}, status=405)

    try:
        url = request.POST.get("url")
        action = request.POST.get("action")

        clean_url = normalize_url(url)
    except Exception as e:
        return JsonResponse({"status": "error", "error": f"Invalid request: {str(e)}"}, status=400)

    # Run the actual scan action
    result = await run_action(action, clean_url)

    # Load or create report
    if await sync_to_async(ScanReport.should_create_new_report)(clean_url):
        report = ScanReport(url=clean_url, results={})
    else:
        report = await sync_to_async(ScanReport.get_latest_report)(clean_url) \
                 or ScanReport(url=clean_url, results={})

    # Update results dict
    results = report.results or {}
    results[action] = result
    report.results = results

    # Save after each action to accumulate results
    if report.pk:
        await sync_to_async(report.save)(update_fields=["results"])
    else:
        await sync_to_async(report.save)()

    return JsonResponse(result)

def download_report(request):
    """Generate and download a PDF security report"""
    url = request.GET.get("url")
    if not url:
        return JsonResponse({"error": "URL parameter is required"}, status=400)
    
    # Normalize the URL
    clean_url = normalize_url(url)
    
    # Get the latest report for this URL
    report = ScanReport.get_latest_report(clean_url)
    if not report:
        return JsonResponse({"error": "No report found for this URL"}, status=404)
    
    # Generate the PDF
    pdf_buffer = generate_security_report_pdf(clean_url, report.results)
    
    # Create the HTTP response with PDF content
    response = HttpResponse(pdf_buffer.read(), content_type='application/pdf')
    safe_url = clean_url.replace("://", "_").replace(".", "_").replace("/", "_")
    response['Content-Disposition'] = f'attachment; filename="security_report_{safe_url}_{datetime.date.today()}.pdf"'
    
    return response



def download_report_docx(request):
    """Generate and download a DOCX security report"""
    url = request.POST.get("url") or request.GET.get("url")
    company_name = request.POST.get("company_name") or request.GET.get("company_name")
    
    if not url:
        return JsonResponse({"error": "URL parameter is required"}, status=400)
    
    clean_url = normalize_url(url)
    report = ScanReport.get_latest_report(clean_url)
    if not report:
        return JsonResponse({"error": "No report found for this URL"}, status=404)
    
    # Fetch page title
    page_title = get_page_title(clean_url)
    logger.info(f"[download_report_docx] Page title for {clean_url}: {page_title}")

    # Extract company name from title if not provided
    if not company_name:
        company_name = get_company_name_from_title(clean_url)
        logger.info(f"[download_report_docx] Extracted company name: {company_name}")
    
    # Generate the DOCX
    docx_buffer = generate_security_report_docx(
        clean_url,
        report.results,
        company_name=company_name,
        page_title=page_title
    )
    
    # Prepare response
    response = HttpResponse(
        docx_buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    
    safe_url = clean_url.replace("://", "_").replace(".", "_").replace("/", "_")
    response['Content-Disposition'] = (
        f'attachment; filename="security_report_{safe_url}_{datetime.date.today()}.docx"'
    )
    
    return response


    """Generate and download a DOCX security report"""
    url = request.POST.get("url") or request.GET.get("url")
    company_name = request.POST.get("company_name") or request.GET.get("company_name")
    
    if not url:
        return JsonResponse({"error": "URL parameter is required"}, status=400)
    
    clean_url = normalize_url(url)
    report = ScanReport.get_latest_report(clean_url)
    if not report:
        return JsonResponse({"error": "No report found for this URL"}, status=404)
    
    # Fetch page title from the website
    page_title = get_page_title(clean_url)
    
    # If no company name provided, try to extract from page title
    if not company_name:
        company_name = get_company_name_from_title(clean_url)
    
    # Generate the DOCX with page title and company name
    docx_buffer = generate_security_report_docx(
        clean_url, 
        report.results, 
        company_name=company_name,
        page_title=page_title
    )
    
    # Create the HTTP response
    response = HttpResponse(
        docx_buffer.read(), 
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    safe_url = clean_url.replace("://", "_").replace(".", "_").replace("/", "_")
    response['Content-Disposition'] = f'attachment; filename="security_report_{safe_url}_{datetime.date.today()}.docx"'
    
    return response
    """Generate and download a DOCX security report"""
    url = request.POST.get("url") or request.GET.get("url")
    if not url:
        return JsonResponse({"error": "URL parameter is required"}, status=400)
    
    clean_url = normalize_url(url)
    report = ScanReport.get_latest_report(clean_url)
    if not report:
        return JsonResponse({"error": "No report found for this URL"}, status=404)
    
    # Generate the DOCX
    docx_buffer = generate_security_report_docx(clean_url, report.results)
    
    # Create the HTTP response
    response = HttpResponse(docx_buffer.read(), 
                          content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    safe_url = clean_url.replace("://", "_").replace(".", "_").replace("/", "_")
    response['Content-Disposition'] = f'attachment; filename="security_report_{safe_url}_{datetime.date.today()}.docx"'
    
    return response
    """Generate readable DOCX security report."""
    url = request.POST.get("url") or request.GET.get("url")
    if not url:
        return JsonResponse({"error": "URL parameter is required"}, status=400)

    clean_url = normalize_url(url)
    report = ScanReport.get_latest_report(clean_url)
    if not report:
        return JsonResponse({"error": "No report found for this URL"}, status=404)

    doc = Document()
    doc.add_heading(f"Website Security Report for {clean_url}", 0)
    doc.add_paragraph(f"Report Date: {datetime.date.today()}").alignment = 1  # center
    doc.add_paragraph("\n")

    # --- Email Security ---
    email_sec = report.results.get('email_security', {}).get('result', {}).get('results', [])
    if email_sec:
        doc.add_heading("Email Security", level=1)
        for item in email_sec:
            check = item.get('check', 'Unknown')
            status = item.get('status', 'Unknown')
            details = item.get('details', '')
            recommendation = item.get('recommendation', '')
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{check}: {status}\nDetails: {details}\nRecommendation: {recommendation}")

    # --- Security Headers ---
    sec_headers = report.results.get('security_headers', {}).get('result', {}).get('results', [])
    if sec_headers:
        doc.add_heading("Security Headers", level=1)
        for header in sec_headers:
            check = header.get('check', 'Unknown')
            status = header.get('status', 'Unknown')
            if status in ('fail', 'missing'):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{check}: {status}\nDetails: {header.get('details', '')}\nRecommendation: {header.get('recommendation', '')}")

    # --- Frontend Libraries ---
    libs = report.results.get('frontend_libraries', {}).get('result', {}).get('results', [])
    if libs:
        doc.add_heading("Frontend Libraries", level=1)
        for lib in libs:
            name = lib.get('name')
            detected = lib.get('detected_version')
            latest = lib.get('latest_version')
            status = lib.get('status')
            rec = lib.get('recommendation', '')
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{name} ({detected} → {latest}): {status}\nRecommendation: {rec}")

    # --- Recommendations ---
    doc.add_heading("Recommendations", level=1)
    general_rec = [
        "Update outdated libraries.",
        "Implement missing security headers.",
        "Follow email security recommendations.",
        "Conduct VAPT testing."
    ]
    for r in general_rec:
        doc.add_paragraph(r, style='List Bullet')

    # Send DOCX
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    safe_url = clean_url.replace("://", "_").replace(".", "_").replace("/", "_")
    response["Content-Disposition"] = f'attachment; filename="security_report_{safe_url}_{datetime.date.today()}.docx"'
    return response
    """Generate and download a DOCX security report similar to PDF formatting"""
    url = request.POST.get("url") or request.GET.get("url")
    if not url:
        return JsonResponse({"error": "URL parameter is required"}, status=400)

    clean_url = normalize_url(url)
    report = ScanReport.get_latest_report(clean_url)
    if not report:
        return JsonResponse({"error": "No report found for this URL"}, status=404)

    doc = Document()
    
    # Title
    title = doc.add_heading(f"Report on Website Security Vulnerabilities for {clean_url}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Date
    doc.add_paragraph(f"Report Date: {datetime.date.today()}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("\n")  # Spacer

    # Intro
    intro_text = (
        f"During our routine monitoring of the {clean_url} domain, "
        "we have identified some security vulnerabilities that require attention."
    )
    doc.add_paragraph(intro_text)

    # Scan details
    doc.add_heading("Preliminary Findings", level=2)

    # Loop through scan results like PDF
    for section, content in report.results.items():
        doc.add_heading(section.replace("_", " ").title(), level=3)
        if isinstance(content, dict):
            for key, value in content.items():
                # Use bulleted list for readability
                p = doc.add_paragraph(f"{key}: {value}", style='List Bullet')
        else:
            doc.add_paragraph(str(content))

    # Recommendations (optional, can mimic PDF bullets)
    doc.add_heading("Recommendations", level=2)
    recommendations = []

    # Example: you can generate recommendations dynamically like PDF
    # Here you just list keys missing
    for section, content in report.results.items():
        if isinstance(content, dict):
            for key, value in content.items():
                recommendations.append(f"Check {key} in {section}")

    for rec in recommendations:
        doc.add_paragraph(rec, style='List Bullet')

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    safe_url = clean_url.replace("://", "_").replace(".", "_").replace("/", "_")
    response["Content-Disposition"] = f'attachment; filename="security_report_{safe_url}_{datetime.date.today()}.docx"'

    return response

