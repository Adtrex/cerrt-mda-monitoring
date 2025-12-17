import io
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_style(doc, text, level=1, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    """Add a heading with custom styling"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = alignment
    
    # Style the heading text
    for run in heading.runs:
        run.font.name = 'Helvetica'
        if level == 0:
            run.font.size = Pt(14)
        elif level == 1:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)
        run.font.bold = True
    
    return heading

def add_paragraph_with_style(doc, text, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, space_after=8):
    """Add a paragraph with custom styling"""
    para = doc.add_paragraph(text)
    para.alignment = alignment
    
    # Style the paragraph
    for run in para.runs:
        run.font.name = 'Helvetica'
        run.font.size = Pt(12)
    
    # Add spacing
    para_format = para.paragraph_format
    para_format.space_after = Pt(space_after)
    
    return para

def add_bullet_point(doc, text):
    """Add a bullet point with custom styling"""
    para = doc.add_paragraph(text, style='List Bullet')
    
    for run in para.runs:
        run.font.name = 'Helvetica'
        run.font.size = Pt(12)
    
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.2
    
    return para

def generate_security_report_docx(url, scan_results, company_name=None, page_title=None):
    """Generate a professional security report in DOCX format
    
    Args:
        url: Target URL that was scanned
        scan_results: Dictionary containing scan results
        company_name: Optional company/agency name to add to header
        page_title: Optional page title from the website (if None, uses domain)
    """
    buffer = io.BytesIO()
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add company name to header if provided
    if company_name:
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = company_name
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in header_para.runs:
            run.font.name = 'Helvetica'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)  # Gray color
        header_para.paragraph_format.space_after = Pt(12)
    
    # Extract domain name
    domain_name = url.replace('https://', '').replace('http://', '').replace('www.', '').split('/')[0]
    
    # Use page title if provided, otherwise use domain name
    display_name = page_title.upper() if page_title else domain_name.upper()
    
    # Date with ordinal suffix
    day = datetime.datetime.now().day
    if 4 <= day <= 20 or 24 <= day <= 30:
        suffix = "th"
    elif day % 10 == 1:
        suffix = "st"
    elif day % 10 == 2:
        suffix = "nd"
    elif day % 10 == 3:
        suffix = "rd"
    else:
        suffix = "th"
    current_date = datetime.datetime.now().strftime(f'%d{suffix} %B %Y')
    
    # Title
    title = add_heading_with_style(
        doc, 
        f"Report on Website Security Vulnerabilities for {display_name}", 
        level=0, 
        alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    )
    
    # Date
    date_para = add_paragraph_with_style(doc, current_date, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=12)
    
    # Add spacer
    doc.add_paragraph()
    
    # Introduction
    intro_text = (
        f"During our routine monitoring of the {display_name} domain ({url}), "
        "we have identified some critical security vulnerabilities that require immediate attention. "
        "These vulnerabilities pose significant risks to the security and privacy of the website's users "
        "and could potentially lead to unauthorized access or the compromise of sensitive information."
    )
    add_paragraph_with_style(doc, intro_text)
    
    # Preliminary Findings
    add_heading_with_style(doc, "Preliminary Findings", level=1)
    
    # --- 1. Outdated Software Packages ---
    if 'frontend_libraries' in scan_results:
        add_heading_with_style(doc, "1. Outdated Software Packages", level=1)
        
        outdated_libs = []
        scan_result = scan_results.get('frontend_libraries', {}).get('result', {}) or {}
        for lib in scan_result.get('results', []):
            if lib.get('status') == 'outdated':
                outdated_libs.append(f"{lib['name']} (Version {lib['detected_version']})")
        
        if outdated_libs:
            lib_text = (
                "We observed the presence of " + 
                ("multiple obsolete software packages" if len(outdated_libs) > 1 else "an obsolete software package") +
                " on the website, these " +
                ("include" if len(outdated_libs) > 1 else "includes") +
                f" {', '.join(outdated_libs)}. " +
                ("These outdated versions are" if len(outdated_libs) > 1 else "This outdated version is") +
                " known to contain vulnerabilities that can be exploited by attackers "
                "to gain unauthorized access to the website or compromise user data."
            )
        else:
            lib_text = (
                "We observed the presence of multiple software packages on the website. "
                "All detected libraries appear to be up to date with no known critical vulnerabilities."
            )
        add_paragraph_with_style(doc, lib_text)
    
    # --- 2. Missing Security Headers ---
    if 'security_headers' in scan_results:
        add_heading_with_style(doc, "2. Missing Security Headers", level=1)
        
        missing_headers = []
        for header in scan_results['security_headers']['result'].get('results', []):
            if header.get('status') in ('fail', 'missing'):
                header_name = header['check']
                if 'x-frame-options' in header_name.lower():
                    missing_headers.append('X-Frame-Options')
                elif 'referrer-policy' in header_name.lower():
                    missing_headers.append('Referrer-Policy')
                elif 'permissions-policy' in header_name.lower():
                    missing_headers.append('Permissions-Policy')
                elif 'content-security-policy' in header_name.lower():
                    missing_headers.append('Content Security Policy')
                elif 'strict-transport-security' in header_name.lower():
                    missing_headers.append('Strict Transport Security')
                elif 'x-content-type-options' in header_name.lower():
                    missing_headers.append('X-Content-Type-Options')
                else:
                    missing_headers.append(header_name)
        
        if missing_headers:
            missing_headers = list(dict.fromkeys(missing_headers))
            header_text = (
                f"Our routine monitoring has revealed that the {display_name} domain is not implementing "
                f"crucial security headers, including {', '.join(missing_headers)}. "
                "These headers play a vital role in safeguarding the website against various types of attacks, "
                "such as cross-site scripting (XSS), clickjacking, and referrer leaks."
            )
            add_paragraph_with_style(doc, header_text)
        else:
            add_paragraph_with_style(doc, "All essential security headers appear to be properly implemented.")
        
        doc.add_paragraph()
    
    # --- 3. Email Authentication ---
    if 'email_security' in scan_results:
        add_heading_with_style(doc, "3. Email Authentication Methods", level=1)
        email_data = scan_results['email_security']
        
        # Normalize structure
        if isinstance(email_data, dict) and 'result' in email_data:
            email_results = email_data['result']
        else:
            email_results = email_data
        
        checks_summary = []
        missing_records = []
        
        if 'results' in email_results and email_results['results']:
            for result in email_results['results']:
                check_name = result.get('check', 'Unknown')
                status = result.get('status', 'fail')
                
                if status == 'pass':
                    checks_summary.append(f"• {check_name} record found")
                else:
                    checks_summary.append(f"• No {check_name} record found")
                    missing_records.append(check_name)
        
        # Add the bulleted checks
        for summary in checks_summary:
            para = doc.add_paragraph(summary)
            para.paragraph_format.left_indent = Inches(0.5)
            for run in para.runs:
                run.font.name = 'Helvetica'
                run.font.size = Pt(12)
        
        # Explanatory text for missing records
        if missing_records:
            explanations = []
            if 'DMARC' in missing_records:
                explanations.append(
                    "The absence of a Domain-based Message Authentication, Reporting and Conformance (DMARC) record "
                    "indicates a lack of specific instructions from the domain owner regarding email handling, "
                    "potentially leaving the domain vulnerable to malicious activities."
                )
            if 'DKIM' in missing_records:
                explanations.append(
                    "The absence of DomainKeys Identified Mail (DKIM) records suggests that the domain owner has not "
                    "implemented this additional layer of authentication. Without DKIM signatures, there is an increased "
                    "risk of email tampering and spoofing, potentially compromising the trustworthiness of communications."
                )
            if 'SPF' in missing_records:
                explanations.append(
                    "The absence of a Sender Policy Framework (SPF) record makes it difficult to prevent spammers "
                    "from forging messages that appear to come from the domain."
                )
            
            add_paragraph_with_style(doc, " ".join(explanations))
    
    # --- Recommendations ---
    add_heading_with_style(doc, "Recommendations", level=1)
    recommendations = []
    
    # Outdated libraries
    if 'frontend_libraries' in scan_results:
        outdated_libs = []
        scan_result = scan_results.get('frontend_libraries', {}).get('result', {}) or {}
        for lib in scan_result.get('results', []):
            if lib.get('status') == 'outdated':
                outdated_libs.append(f"{lib['name']} (Version {lib['detected_version']})")
        
        if outdated_libs:
            recommendations.append(
                "It is strongly recommended that the identified software packages be updated to their latest versions. "
                "This action will help mitigate the risk associated with known vulnerabilities and ensure the overall "
                "security of your website."
            )
    
    # Missing security headers
    if 'security_headers' in scan_results:
        missing_headers = []
        for header in scan_results['security_headers']['result'].get('results', []):
            if header.get('status') in ('fail', 'missing'):
                header_name = header['check']
                if 'x-frame-options' in header_name.lower():
                    missing_headers.append('X-Frame-Options')
                elif 'referrer-policy' in header_name.lower():
                    missing_headers.append('Referrer-Policy')
                elif 'permissions-policy' in header_name.lower():
                    missing_headers.append('Permissions-Policy')
                elif 'content-security-policy' in header_name.lower():
                    missing_headers.append('Content Security Policy')
                elif 'strict-transport-security' in header_name.lower():
                    missing_headers.append('Strict Transport Security')
                elif 'x-content-type-options' in header_name.lower():
                    missing_headers.append('X-Content-Type-Options')
                else:
                    missing_headers.append(header_name)
        
        if missing_headers:
            missing_headers = list(dict.fromkeys(missing_headers))
            recommendations.append(
                f"Implement the missing security headers ({', '.join(missing_headers)}) on the {display_name} domain. "
                "These headers are vital in protecting the site against XSS, clickjacking, and data leakage attacks."
            )
    
    # Missing email records
    if 'email_security' in scan_results:
        email_data = scan_results['email_security']
        if isinstance(email_data, dict) and 'result' in email_data:
            email_results = email_data['result']
        else:
            email_results = email_data
        
        missing_records = []
        if 'results' in email_results and email_results['results']:
            for result in email_results['results']:
                check_name = result.get('check', 'Unknown')
                status = result.get('status', 'fail')
                if status != 'pass':
                    missing_records.append(check_name)
        
        if "SPF" in missing_records:
            recommendations.append("Add an SPF record to prevent spammers from forging emails from your domain.")
        if "DKIM" in missing_records:
            recommendations.append("Add a DKIM record to ensure email integrity and reduce spoofing risks.")
        if "DMARC" in missing_records:
            recommendations.append("Add a DMARC record to provide email handling instructions and protect against phishing.")
    
    # Always add general recommendation
    recommendations.append(
        "Conduct a Vulnerability Assessment and Penetration Testing (VAPT) to identify and address additional security weaknesses."
    )
    
    # Render recommendations as bullet points
    for rec in recommendations:
        add_bullet_point(doc, rec)
    
    # Conclusion
    add_heading_with_style(doc, "Conclusion", level=1)
    conclusion = (
        "Addressing the identified vulnerabilities is essential to mitigate potential security risks "
        "and uphold trust and integrity of your website. We strongly recommend you to take immediate action "
        "to implement the recommended measures."
    )
    add_paragraph_with_style(doc, conclusion)
    
    # Save to buffer
    doc.save(buffer)
    buffer.seek(0)
    return buffer